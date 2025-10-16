import fitz
import pytesseract
from PIL import Image
import io
import asyncio
from pathlib import Path
from docx import Document as DocxDocument
import tempfile
import subprocess
from .base import *
from ...exception import CustomException
import sys

class DOCXParser(BaseParser):
    
    def can_handle(self, file_type: str) -> bool:
        return file_type == "docx"
    
    async def parse(self, path: Path) -> str:
        """Parse DOCX with native-first, OCR fallback strategy"""
        try:
            has_readable_content = await self._detect_readable_docx(path)
            
            if has_readable_content:
                return await self._extract_native_docx(path)
            else:
                return await self._extract_docx_with_ocr(path)
                
        except Exception as e:
            if not isinstance(e, CustomException):
                e = CustomException(e, sys)
            self.logger.error(e)
            raise e
    
    async def _detect_readable_docx(self, file_path: Path) -> bool:
        """Check if DOCX has readable text content"""
        loop = asyncio.get_event_loop()
        
        def _check_content():
            try:
                doc = DocxDocument(file_path)
                total_chars = 0
                
                # Sample first few paragraphs
                for i, para in enumerate(doc.paragraphs):
                    if i >= 5:  # Check first 5 paragraphs
                        break
                    total_chars += len(para.text.strip())
                
                # If decent amount of text found, use native
                return total_chars > 50
            except Exception:
                return False
        
        return await loop.run_in_executor(None, _check_content)
    
    async def _extract_native_docx(self, file_path: Path) -> str:
        """Extract using python-docx library"""
        loop = asyncio.get_event_loop()
        
        def _extract():
            full_text = ""
            doc = DocxDocument(file_path)
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text += para.text.strip() + "\n"
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                full_text += f"\nTable {i}\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        full_text += row_text + "\n"
            
            return full_text
        
        result = await loop.run_in_executor(None, _extract)
        self.logger.info("parsed through native method")
        return result
    
    async def _ocr_main(self, page, dpi: int = 400) -> str:
        loop = asyncio.get_event_loop()
        
        def _do_ocr():
            pix = page.get_pixmap(dpi=dpi)
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            ocr_result = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
            return " ".join([t for t in ocr_result['text'] if t.strip()])
        
        return await loop.run_in_executor(None, _do_ocr)
    
    async def _extract_docx_with_ocr(self, file_path: Path) -> str:
        """Fallback: Convert DOCX to images and OCR"""
        loop = asyncio.get_event_loop()
        full_text = ""
        
        def _convert_to_pdf():
            try:
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                    temp_pdf_path = Path(temp_pdf.name)
                    
                    # Convert using libreoffice
                    result = subprocess.run([
                        'libreoffice', '--headless', '--convert-to', 'pdf',
                        '--outdir', str(temp_pdf_path.parent), str(file_path)
                    ], capture_output=True, timeout=60)
                    
                    if result.returncode == 0:
                        # Find the generated PDF
                        pdf_name = file_path.stem + '.pdf'
                        actual_pdf_path = temp_pdf_path.parent / pdf_name
                        return actual_pdf_path if actual_pdf_path.exists() else None
                    return None
            except Exception as e:
                self.logger.error(f"DOCX to PDF conversion failed: {e}")
                return None
        
        try:
            pdf_path = await loop.run_in_executor(None, _convert_to_pdf)
            
            if pdf_path and pdf_path.exists():
                # OCR the PDF
                pdf_doc = fitz.open(str(pdf_path))
                
                for page in pdf_doc:
                    page_text = await self._ocr_main(page)
                    if len(page_text) < 50:
                        page_text = await self._ocr_main(page, dpi=600)
                    full_text += page_text + "\n"
                
                pdf_doc.close()
                pdf_path.unlink()  # Clean up
                self.logger.info("parsed through ocr")
            else:
                raise Exception("Failed to convert DOCX to PDF")
                
        except Exception as e:
            e = CustomException(e, sys)
            self.logger.error(e)
            # Last resort - try basic extraction
            try:
                result = await self._extract_native_docx(file_path)
                return result
            except:
                raise e
        
        return full_text

__all__ = ["DOCXParser"]
